
import os
import socketio as socketio_lib
import io
import re
import json
import math
import uuid
import base64
import random
import asyncio
import tempfile
from datetime import datetime
from typing import List, Dict, Optional
from io import BytesIO

import requests
import httpx
import aiohttp
import soundfile as sf
import numpy as np
import librosa
import noisereduce as nr
from pydub import AudioSegment, effects
from docx import Document
from openpyxl import Workbook
from pptx import Presentation
from pptx.util import Inches, Pt
from pptx.dml.color import RGBColor
from pypdf import PdfReader
from fastapi import FastAPI, UploadFile, HTTPException, Body, Form, APIRouter, Query, Request
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from openai import OpenAI
from jose import jwt, JWTError

import firebase_admin
from firebase_admin import credentials, storage, firestore
import os, sys, logging
import socketio as socketio_lib

# Import error handler
from core.error_handler import setup_error_handlers, CustomHTTPException, ValidationError, BusinessLogicError, ExternalServiceError
from core.language_support import (
    build_ai_detection_messages,
    extract_generator_info,
    format_ai_detection_summary,
    normalize_language,
    nsfw_flag_from_value,
    quality_flag_from_value,
)
# NOTE: image_edit_router temporarily disabled (frontend handles messaging)
from endpoints.chat import router as chat_router
from endpoints.chat.first_prompt import router as chat_first_prompt_router
from endpoints.chat.manage import router as chat_manage_router
from endpoints.agent import router as agent_router
from endpoints.chat_title import router as chat_title_router
from endpoints.presentation import router as presentation_router
from endpoints.generate_image.gemini_image import router as gemini_image_router
from endpoints.video_gemini.gemini_video import router as gemini_video_router
from endpoints.generate_image.generateImageGeminiSearch import router as gemini_image_search_router
from endpoints.generate_image.edit_image_gemini import router as gemini_image_edit_router
from endpoints.generate_image.analyze_image_gemini import router as gemini_image_analyze_router
from endpoints.deep_research import router as deep_research_router
from endpoints.web_search import router as web_search_router
from endpoints.web_link import router as web_link_router
from endpoints.social_posts import router as social_posts_router
from endpoints.generate_doffice.generate_doc import router as generate_doc_router
from endpoints.generate_doffice.generate_ppt import router as generate_ppt_router
from endpoints.generate_doffice.generate_pdf import router as generate_pdf_router
from endpoints.files_word import (
    word_summary_router,
    word_analyze_router,
    word_qna_router,
    word_translate_router,
    word_rewrite_router,
    word_compare_router,
    word_extract_router,
    word_classify_router,
    word_multi_analyze_router,
    word_ocr_extract_router,
    word_layout_router,
    word_deep_extract_router,
    word_grounded_search_router,
    word_structure_export_router,
)
from endpoints.files_pptx import (
    pptx_summary_router,
    pptx_analyze_router,
    pptx_qna_router,
    pptx_translate_router,
    pptx_rewrite_router,
    pptx_compare_router,
    pptx_deep_extract_router,
    pptx_grounded_search_router,
    pptx_structure_export_router,
)
from endpoints.stt_and_tts import stt_router, tts_router
from endpoints.file_export import router as export_chat_router
from endpoints.files_pdf import (
    pdf_analyze_router,
    pdf_summary_router,
    pdf_qna_router,
    pdf_extract_router,
    pdf_compare_router,
    pdf_rewrite_router,
    pdf_classify_router,
    pdf_multianalyze_router,
    pdf_ocr_extract_router,
    pdf_layout_router,
    pdf_deepextract_router,
    pdf_grounded_search_router,
    pdf_translate_router,
    pdf_structure_export_router,
    pdf_test_router,
)
# from endpoints.image_edit import router as image_edit_router
from core.websocket_manager import sio


router = APIRouter()

# --- LOGGING FORCE CONFIG (put this at the very top of main.py) ---


LOG_LEVEL = os.getenv("LOG_LEVEL", "DEBUG").upper()
LOG_FORMAT = os.getenv("LOG_FORMAT", "%(asctime)s %(levelname)s [%(name)s] %(message)s")

# Root logger -> stdout, force override uvicorn defaults
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.DEBUG),
    format=LOG_FORMAT,
    handlers=[logging.StreamHandler(sys.stdout)],
    force=True,  # override any existing handlers from uvicorn etc.
)

root = logging.getLogger()

# Make uvicorn loggers use the same handlers/level (error + access + general)
for name in ("uvicorn", "uvicorn.error", "uvicorn.access"):
    lg = logging.getLogger(name)
    lg.handlers = root.handlers
    lg.setLevel(getattr(logging, LOG_LEVEL, logging.DEBUG))
    lg.propagate = False

# (İsteğe bağlı) Sık kullanılan kütüphaneleri de aynı seviyeye çek
for name in (
    "httpx", "asyncio", "pydantic", "pdf2image", "pytesseract",
    "pypdf", "pptx", "docx", "PIL"
):
    lg = logging.getLogger(name)
    lg.handlers = root.handlers
    lg.setLevel(getattr(logging, LOG_LEVEL, logging.DEBUG))
    lg.propagate = False

# Render/containers: anında flush için
try:
    sys.stdout.reconfigure(line_buffering=True)  # Py3.7+
except Exception:
    pass
# --- END LOGGING FORCE CONFIG ---



from dotenv import load_dotenv
load_dotenv()

FIREBASE_SERVICE_ACCOUNT_BASE64 = os.getenv("FIREBASE_SERVICE_ACCOUNT_BASE64")

# Make Firebase init optional for local runs without credentials
service_account_info = None
if FIREBASE_SERVICE_ACCOUNT_BASE64:
    try:
        decoded_json = base64.b64decode(FIREBASE_SERVICE_ACCOUNT_BASE64).decode('utf-8')
        service_account_info = json.loads(decoded_json)
        if not firebase_admin._apps:
            cred = credentials.Certificate(service_account_info)
            firebase_admin.initialize_app(cred, {
                'storageBucket': 'aveniaapp.firebasestorage.app'
            })
    except Exception as exc:
        logging.getLogger("pdfread.firebase").warning(
            "Failed to initialize Firebase from FIREBASE_SERVICE_ACCOUNT_BASE64: %s", exc
        )
else:
    logging.getLogger("pdfread.firebase").info(
        "FIREBASE_SERVICE_ACCOUNT_BASE64 not set; skipping Firebase initialization"
    )


app = FastAPI(title="Avenia PDF Read API", version="1.0.0")

# Setup error handlers
setup_error_handlers(app)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(chat_router)
app.include_router(chat_first_prompt_router)
app.include_router(chat_manage_router)
app.include_router(agent_router)
app.include_router(chat_title_router)
app.include_router(deep_research_router)
app.include_router(web_search_router)
app.include_router(web_link_router)
app.include_router(social_posts_router)
app.include_router(presentation_router)
app.include_router(gemini_image_router)
app.include_router(gemini_video_router)
app.include_router(gemini_image_search_router)
app.include_router(gemini_image_edit_router)
app.include_router(gemini_image_analyze_router)
# Office doc generation
app.include_router(generate_doc_router)
app.include_router(generate_ppt_router)
app.include_router(generate_pdf_router)
# Word/PPTX
app.include_router(word_summary_router)
app.include_router(word_analyze_router)
app.include_router(word_qna_router)
app.include_router(word_translate_router)
app.include_router(word_rewrite_router)
app.include_router(word_compare_router)
app.include_router(word_extract_router)
app.include_router(word_classify_router)
app.include_router(word_multi_analyze_router)
app.include_router(word_ocr_extract_router)
app.include_router(word_layout_router)
app.include_router(word_deep_extract_router)
app.include_router(word_grounded_search_router)
app.include_router(word_structure_export_router)
app.include_router(pptx_summary_router)
app.include_router(pptx_analyze_router)
app.include_router(pptx_qna_router)
app.include_router(pptx_translate_router)
app.include_router(pptx_rewrite_router)
app.include_router(pptx_compare_router)
app.include_router(pptx_deep_extract_router)
app.include_router(pptx_grounded_search_router)
app.include_router(pptx_structure_export_router)
# PDF document processing
app.include_router(pdf_analyze_router)
app.include_router(pdf_summary_router)
app.include_router(pdf_qna_router)
app.include_router(pdf_extract_router)
app.include_router(pdf_compare_router)
app.include_router(pdf_rewrite_router)
app.include_router(pdf_classify_router)
app.include_router(pdf_multianalyze_router)
app.include_router(pdf_ocr_extract_router)
app.include_router(pdf_layout_router)
app.include_router(pdf_deepextract_router)
app.include_router(pdf_grounded_search_router)
app.include_router(pdf_translate_router)
app.include_router(pdf_structure_export_router)
app.include_router(pdf_test_router)
# app.include_router(image_edit_router)
app.include_router(stt_router)
app.include_router(tts_router)
app.include_router(export_chat_router)

socket_app = socketio_lib.ASGIApp(sio, other_asgi_app=app)

socket_app = socketio_lib.ASGIApp(sio, other_asgi_app=app)

JWT_SECRET = os.getenv("JWT_HS_SECRET", "change_me_in_production")
JWT_ISSUER = os.getenv("JWT_ISS", "chatgbtmini")
JWT_AUDIENCE = os.getenv("JWT_AUD", "chatgbtmini-mobile")

PUBLIC_PATHS = {"/healthz", "/health", "/docs", "/openapi.json", "/redoc"}

auth_logger = logging.getLogger("pdfread.auth")

from fastapi.openapi.utils import get_openapi  # noqa: E402


def custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    openapi_schema = get_openapi(
        title=app.title,
        version=app.version,
        description="Avenia PDF Read API",
        routes=app.routes,
    )
    openapi_schema.setdefault("components", {}).setdefault("securitySchemes", {})[
        "BearerAuth"
    ] = {
        "type": "http",
        "scheme": "bearer",
        "bearerFormat": "JWT",
    }
    openapi_schema["security"] = [{"BearerAuth": []}]
    app.openapi_schema = openapi_schema
    return app.openapi_schema


app.openapi = custom_openapi


def _verify_bearer_token(auth_header: str) -> Dict:
    if not auth_header:
        auth_logger.warning("Authorization header missing")
        raise HTTPException(
            status_code=401,
            detail={"error": "access_denied", "message": "Access token required"},
        )

    parts = auth_header.strip().split(" ")
    if len(parts) != 2 or parts[0].lower() != "bearer":
        auth_logger.warning("Invalid authorization header format")
        raise HTTPException(
            status_code=401,
            detail={"error": "invalid_token", "message": "Invalid authorization header"},
        )

    token = parts[1]
    try:
        payload = jwt.decode(
            token,
            JWT_SECRET,
            algorithms=["HS256"],
            audience=JWT_AUDIENCE,
            issuer=JWT_ISSUER,
        )
        auth_logger.debug("Token verified for subject: %s", payload.get("sub"))
        return payload
    except JWTError as exc:
        auth_logger.warning("Token verification failed: %s", exc)
        raise HTTPException(
            status_code=401,
            detail={"error": "invalid_token", "message": "Invalid or expired access token"},
        )
    except Exception as exc:  # pragma: no cover - defensive
        auth_logger.error("Unexpected error while verifying token: %s", exc)
        raise HTTPException(
            status_code=500,
            detail={"error": "auth_error", "message": "Failed to validate access token"},
        )


@app.middleware("http")
async def authenticate_request(request: Request, call_next):
    path = request.url.path
    if path.startswith("/static/") or path in PUBLIC_PATHS or request.method == "OPTIONS":
        return await call_next(request)

    try:
        payload = _verify_bearer_token(request.headers.get("Authorization"))
        request.state.token_payload = payload
    except HTTPException as exc:
        detail = exc.detail if isinstance(exc.detail, dict) else {"detail": exc.detail}
        return JSONResponse(status_code=exc.status_code, content=detail)

    return await call_next(request)

DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-4")
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
RUNWAY_API_KEY = os.getenv("RUNWAY_API_KEY")
ELEVENLABS_API_KEY = os.getenv("ELEVENLABS_API_KEY")
TTS_VOICE_ID = os.getenv("ELEVENLABS_VOICE_ID", "21m00Tcm4TlvDq8ikWAM")  # Default voice
API_KEY = os.getenv("AIORNOT_API_KEY")  # aiornot API key


class DocRequest(BaseModel):
    prompt: str

async def wait_for_video_ready(video_id, retries=30, delay=5):
    status_url = f"https://api.dev.runwayml.com/v1/tasks/{video_id}"  # ✅ doğru endpoint
    headers = {
        "Authorization": f"Bearer {RUNWAY_API_KEY}",
        "Content-Type": "application/json",
        "X-Runway-Version": "2024-11-06"
    }

    async with httpx.AsyncClient() as client:
        for attempt in range(retries):
            print(f"🎞️ Video durumu sorgulanıyor (deneme {attempt + 1})...")
            response = await client.get(status_url, headers=headers)

            try:
                data = response.json()
            except Exception:
                print(f"❌ JSON çözümleme hatası (deneme {attempt + 1}):", response.text)
                print(response.status_code, response.text) 
                await asyncio.sleep(delay)
                continue

            status = data.get("status")
            print(f"📌 Durum: {status}")

            if status == "SUCCEEDED":
                print("🧾 API'den dönen tüm veri (debug):")
                print("------------------------")
                print(json.dumps(data, indent=2))
                print("------------------------")
                video_output = data.get("output")
                if not video_output or not isinstance(video_output,list) or not video_output[0]:
                    raise Exception("✅ Video üretildi ama videoUri bulunamadı.")
                return video_output[0]

            if status == "FAILED":
                error_detail = data.get("error", {}).get("message", "Bilinmeyen hata")
                raise HTTPException(
                    status_code=422,  # Unprocessable Entity
                    detail={
                        "type": "runway_failure",
                        "message": f"Runway video üretimi başarısız oldu: {error_detail}",
                        "runway_response": data
                    }
                )

            await asyncio.sleep(delay)

    raise Exception("⚠️ Video üretimi zaman aşımına uğradı.")


def extract_text_from_pdf(path: str) -> str:
    print("[extract_text_from_pdf] 📥 Dosya okunuyor:", path)
    reader = PdfReader(path)
    all_text = ""
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        print(f"[extract_text_from_pdf] 📄 Sayfa {i+1} okundu, karakter:", len(page_text) if page_text else 0)
        if page_text:
            all_text += page_text + "\n"

    print("[extract_text_from_pdf] 🧾 Toplam metin uzunluğu:", len(all_text))
    print("[extract_text_from_pdf] 🔍 İlk 500 karakter:\n", all_text[:500])

    return all_text[:4000]

def ask_gpt_summary(text: str) -> str:
    print("[ask_gpt_summary] 🤖 GPT ile özetleme başlıyor...")
    prompt = f"Bu PDF dosyasının içeriğini kısaca özetle:\n\n{text}"
    print("[ask_gpt_summary] 📤 Gönderilen prompt uzunluğu:", len(prompt))
    print("[ask_gpt_summary] 🔍 İlk 500 karakter:\n", prompt[:500])

    response = client.chat.completions.create(
        model=DEFAULT_MODEL,
        messages=[
            {"role": "system", "content": "Sen profesyonel bir özetleyicisin."},
            {"role": "user", "content": prompt},
        ],
    )
    result = response.choices[0].message.content.strip()
    print("[ask_gpt_summary] ✅ GPT özeti alındı:\n", result)
    return result

# PDF metni üzerinden soru cevap
from fastapi import FastAPI, Form, HTTPException
from fastapi.responses import JSONResponse
from openai import OpenAI
import os


def parse_ppt_prompt(text: str):
    slides = []
    current_slide = {"title": "", "content": "", "image": ""}

    for line in text.splitlines():
        line = line.strip()
        if line.lower().startswith("# slide"):
            if current_slide["title"] or current_slide["content"]:
                slides.append(current_slide)
            current_slide = {"title": "", "content": "", "image": ""}
        elif line.lower().startswith("title:"):
            current_slide["title"] = line.split(":", 1)[1].strip()
        elif line.lower().startswith("content:"):
            current_slide["content"] = line.split(":", 1)[1].strip()
        elif line.lower().startswith("image:"):
            current_slide["image"] = line.split(":", 1)[1].strip()

    if current_slide["title"] or current_slide["content"]:
        slides.append(current_slide)

    return slides


# Firebase init
# 🔊 Yeni endpoint: Speech-to-Text
# ----------------------------------------------------
# TTS: Chat geçmişini sese çevir (OpenAI tts-1)
# ----------------------------------------------------
# ----------------------------------------------------
# Audio Isolation (lokal): hafif denoise + filtreleme
# ----------------------------------------------------
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text_parts = []

    # Paragraflar
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Tablolar
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)



    # ----------------------------
# EMBEDDINGS DESTEKLİ ARAMA
# ----------------------------


# 1. Embedding oluşturma
def create_embedding(text: str) -> List[float]:
    response = client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return response.data[0].embedding

# 2. Cosine similarity
def cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    dot = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = math.sqrt(sum(a * a for a in vec1))
    norm2 = math.sqrt(sum(b * b for b in vec2))
    return dot / (norm1 * norm2)

# 3. PDF yüklenince embedding kaydetme (mevcut summarize_pdf_url sonrası çağır)
def save_embeddings_to_firebase(user_id: str, chat_id: str, file_id: str, file_text: str, summary: str, file_type: str):
    db = firestore.client()

    print(f"[save_embeddings_to_firebase] 📥 Başlatıldı")
    print(f"   → user_id: {user_id}")
    print(f"   → chat_id: {chat_id}")
    print(f"   → file_id: {file_id}")
    print(f"   → file_type: {file_type}")
    print(f"   → file_text uzunluğu: {len(file_text) if file_text else 0}")
    print(f"   → summary uzunluğu: {len(summary) if summary else 0}")

    # Base reference: embeddings/userId/chatId
    base_ref = db.collection("embeddings").document(user_id).collection(chat_id)
    print(f"[save_embeddings_to_firebase] 🔗 Firestore path: embeddings/{user_id}/{chat_id}")

    # 1. Meta veriyi kaydet (chunk_index = -999)
    try:
        meta_doc = {
            "file_id": file_id,
            "chunk_index": -999,
            "summary": summary,
            "type": file_type,
            "createdAt": firestore.SERVER_TIMESTAMP
        }
        base_ref.add(meta_doc)
        print(f"[save_embeddings_to_firebase] ✅ Meta veri kaydedildi: {meta_doc}")
    except Exception as e:
        print(f"[save_embeddings_to_firebase] ❌ Meta kaydı hatası: {e}")

    # 2. Summary’nin embedding’i
    if summary and summary.strip():
        try:
            summary_clean = " ".join(summary.split())
            print(f"[save_embeddings_to_firebase] 🔄 Summary temizlendi: {summary_clean[:100]}...")
            summary_embedding = create_embedding(summary_clean)
            print(f"[save_embeddings_to_firebase] 🧠 Summary embedding boyutu: {len(summary_embedding)}")

            base_ref.add({
                "file_id": file_id,
                "chunk_index": -1,
                "text": summary_clean,
                "embedding": summary_embedding
            })
            print(f"[save_embeddings_to_firebase] ✅ Summary embedding kaydedildi (chunk_index=-1)")
        except Exception as e:
            print(f"[save_embeddings_to_firebase] ❌ Summary embedding hatası: {e}")

    # 3. Chunk embedding’leri
    chunks = [file_text[i:i+500] for i in range(0, len(file_text), 500)]
    print(f"[save_embeddings_to_firebase] 🔄 Toplam {len(chunks)} chunk üretildi")

    for idx, chunk in enumerate(chunks):
        try:
            chunk_clean = " ".join(chunk.split())
            print(f"[save_embeddings_to_firebase] → Chunk {idx} temizlendi (ilk 80 karakter): {chunk_clean[:80]}")

            embedding = create_embedding(chunk_clean)
            print(f"[save_embeddings_to_firebase] 🧠 Chunk {idx} embedding boyutu: {len(embedding)}")

            base_ref.add({
                "file_id": file_id,
                "chunk_index": idx,
                "text": chunk_clean,
                "embedding": embedding
            })
            print(f"[save_embeddings_to_firebase] ✅ Chunk {idx} kaydedildi")
        except Exception as e:
            print(f"[save_embeddings_to_firebase] ❌ Chunk {idx} hatası: {e}")

    print("[save_embeddings_to_firebase] 🎉 Tüm kayıt işlemleri tamamlandı.")


# 4. Yeni endpoint: embeddings tabanlı soru-cevap
##
IMAGE_ENDPOINT = "https://api.aiornot.com/v1/reports/image"
MOCK_MODE = os.getenv("MOCK_MODE", "0") == "1"

# --- Firebase init (service account JSON'u base64 ile ENV'den) ---
FIREBASE_SERVICE_ACCOUNT_BASE64 = os.getenv("FIREBASE_SERVICE_ACCOUNT_BASE64")
if FIREBASE_SERVICE_ACCOUNT_BASE64 and not firebase_admin._apps:
    try:
        decoded = base64.b64decode(FIREBASE_SERVICE_ACCOUNT_BASE64).decode("utf-8")
        cred = credentials.Certificate(json.loads(decoded))
        firebase_admin.initialize_app(cred)
    except Exception as exc:
        logging.getLogger("pdfread.firebase").warning(
            "Deferred Firebase init failed: %s", exc
        )

# Create Firestore client only if Firebase is initialized
db = firestore.client() if firebase_admin._apps else None
##

def decode_base64_maybe_data_url(s: str) -> bytes:
    """Hem düz base64'ü hem de data URL'yi destekler."""
    if not isinstance(s, str):
        raise ValueError("image_base64 must be a string")
    if s.startswith("data:"):
        m = re.match(r"^data:[^;]+;base64,(.+)$", s)
        if not m:
            raise ValueError("Invalid data URL")
        s = m.group(1)
    return base64.b64decode(s)

def interpret_messages_legacy(data, language: Optional[str] = None):
    """Generates localized message list."""
    report = data.get("report", {}) or {}
    facets = data.get("facets", {}) or {}

    verdict = report.get("verdict")
    ai_conf = float((report.get("ai", {}) or {}).get("confidence", 0.0) or 0.0)
    human_conf = float((report.get("human", {}) or {}).get("confidence", 0.0) or 0.0)

    quality_flag = quality_flag_from_value(facets.get("quality"))
    nsfw_flag = nsfw_flag_from_value(facets.get("nsfw"))

    generator_name, generator_conf = extract_generator_info(report.get("generator"))
    if generator_conf is not None and generator_conf < 0.7:
        generator_name, generator_conf = None, None

    return build_ai_detection_messages(
        verdict,
        ai_conf,
        human_conf,
        quality_flag,
        nsfw_flag,
        language=language,
        generator_name=generator_name,
        generator_confidence=generator_conf,
    )


def format_summary_tr(data, language: Optional[str] = None, subject: str = "image") -> str:
    """Returns localized summary for AI detection result."""
    report = data.get("report", {}) or {}
    facets = data.get("facets", {}) or {}

    verdict = report.get("verdict")
    ai_conf = float((report.get("ai", {}) or {}).get("confidence", 0.0) or 0.0)
    human_conf = float((report.get("human", {}) or {}).get("confidence", 0.0) or 0.0)

    quality_flag = quality_flag_from_value(facets.get("quality"))
    nsfw_flag = nsfw_flag_from_value(facets.get("nsfw"))

    return format_ai_detection_summary(
        verdict,
        ai_conf,
        human_conf,
        quality_flag,
        nsfw_flag,
        language=language,
        subject=subject,
    )


def _save_asst_message(user_id: str, chat_id: str, content: str, raw: dict, language: str = "tr"):
    """Writes assistant message to Firestore if initialized."""
    if db is None:
        print("[/analyze-image] Firebase not initialized; skipping Firestore write")
        return None
    try:
        messages_ref = db.collection("users").document(user_id) \
            .collection("chats").document(chat_id) \
            .collection("messages")
        doc_ref = messages_ref.add({
            "role": "assistant",
            "content": content,
            "meta": {
                "language": normalize_language(language),
                "ai_detect": {
                    "verdict": raw.get("report", {}).get("verdict"),
                    "ai_confidence": (raw.get("report", {}).get("ai", {}) or {}).get("confidence"),
                    "human_confidence": (raw.get("report", {}).get("human", {}) or {}).get("confidence"),
                    "quality": (raw.get("facets", {}).get("quality", {}) or {}),
                    "nsfw": (raw.get("facets", {}).get("nsfw", {}) or {}),
                    "raw": raw
                }
            },
            "timestamp": datetime.utcnow(),
            "createdAt": datetime.utcnow()
        })
        message_id = doc_ref[1].id if isinstance(doc_ref, tuple) else doc_ref.id
        print("[/analyze-image] Firestore kaydı:", message_id)
        return {"message_id": message_id, "path": f"users/{user_id}/chats/{chat_id}/messages"}
    except Exception as e:
        print("[/analyze-image] Firestore yazım hatası:", e)
        return None

import endpoints.generate_doffice.generate_video
import endpoints.generate_doffice.generate_video_prompt
import endpoints.generate_doffice.generate_excel
from endpoints.stt_and_tts import stt_router, tts_router
import endpoints.search_docs
import endpoints.healthz
import endpoints.ai_or_not.ai_analyze_image
import endpoints.generate_doffice.generate_doc_advanced
import endpoints.generate_doffice.generate_ppt_advanced
import endpoints.ai_or_not.ai_detect_video
try:
    import endpoints.ai_or_not.check_ai
except Exception as exc:
    logging.getLogger("pdfread.endpoints").warning(
        "check_ai endpoint disabled (optional deps missing): %s", exc
    )
import endpoints.convert_office.pdf_to_word
import endpoints.convert_office.pdf_to_ppt
import endpoints.convert_office.pdf_to_excel
import endpoints.convert_office.word_to_pdf
import endpoints.convert_office.ppt_to_pdf
import endpoints.convert_office.excel_to_pdf


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))
    app.run(host="0.0.0.0", port=port)
