import os
import io
import uuid
import tempfile
import logging
import hashlib  # sadece log için
from fastapi import UploadFile, HTTPException
from main import app, storage
from docx import Document
from fpdf import FPDF

logger = logging.getLogger(__name__)


def _find_unicode_font() -> str | None:
    """
    Ortam değişkeni kullanmadan, yaygın lokasyonlarda bir Unicode TTF arar.
    Reponuza 'fonts/DejaVuSans.ttf' eklerseniz ilk iki yol zaten vurur.
    """
    here = os.path.dirname(__file__)
    candidates = [
        os.path.join(here, "DejaVuSans.ttf"),
        os.path.join(here, "fonts", "DejaVuSans.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "C:\\Windows\\Fonts\\arialuni.ttf",  # Windows
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",  # macOS
    ]
    for p in candidates:
        if p and os.path.exists(p):
            return p
    return None


@app.post("/word-to-pdf")
async def word_to_pdf(file: UploadFile):
    """Convert an uploaded DOCX file to PDF and return a Firebase URL."""
    try:
        logger.info("[/word-to-pdf] START filename=%s content_type=%s",
                    getattr(file, "filename", None), getattr(file, "content_type", None))

        # 1) Dosyayı oku
        doc_bytes = await file.read()
        logger.info("[/word-to-pdf] file.read ok bytes=%d md5=%s",
                    len(doc_bytes), hashlib.md5(doc_bytes).hexdigest())

        # 2) DOCX'i yükle
        document = Document(io.BytesIO(doc_bytes))
        para_count = len(document.paragraphs)
        logger.info("[/word-to-pdf] DOCX parsed paragraphs=%d", para_count)
        if para_count:
            first_preview = (document.paragraphs[0].text or "")[:120]
            logger.debug("[/word-to-pdf] first paragraph preview=%r", first_preview)

        # 3) PDF kurulum
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # -- Unicode font dene (ENV YOK) — bulunamazsa core fonta düş
        ufont = _find_unicode_font()
        if ufont:
            try:
                pdf.add_font("Uni", "", ufont, uni=True)
                pdf.set_font("Uni", size=12)
                logger.info("[/word-to-pdf] font=Uni:12 path=%s", ufont)
            except Exception as fe:
                logger.warning("[/word-to-pdf] Unicode font load failed (%s). Fallback to Arial/latin-1 replace.", fe)
                pdf.set_font("Arial", size=12)
        else:
            pdf.set_font("Arial", size=12)
            logger.info("[/word-to-pdf] unicode font not found → font=Arial:12 (core)")

        # 4) Paragrafları yaz (Unicode yoksa bile 500 atmasın diye güvenli fallback)
        written = 0
        eff_w = getattr(pdf, "epw", pdf.w - pdf.l_margin - pdf.r_margin)  # efektif sayfa genişliği
        logger.info(
            "[/word-to-pdf] write loop start eff_w=%.2f l_margin=%.2f r_margin=%.2f x=%.2f",
            eff_w, pdf.l_margin, pdf.r_margin, pdf.get_x()
        )

        for para in document.paragraphs:
            raw = para.text or ""
            # kontrol karakterlerini normalize et
            text = raw.replace("\r", "").replace("\t", "    ")

            # boş satırsa sadece aşağı in
            if not text.strip():
                pdf.ln(5)
                written += 1
                continue

            try:
                # her satır öncesi X'i sol marja al ve sabit genişlik ver
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(eff_w, 10, text)
            except Exception as we:
                logger.debug(
                    "[/word-to-pdf] multicell error (%s) → hard-wrap ile yeniden dene", we
                )
                # Çok uzun tek kelimeleri kırmak için 60 karakterde bir satır sonu ekle
                if len(text) > 60 and " " not in text[:80]:
                    chunks = [text[i:i+60] for i in range(0, len(text), 60)]
                    text_hw = "\n".join(chunks)
                else:
                    # kelime araları var ama yine de patladıysa yine de kır
                    text_hw = "\n".join([text[i:i+120] for i in range(0, len(text), 120)])

                try:
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(eff_w, 10, text_hw)
                except Exception as we2:
                    logger.debug(
                        "[/word-to-pdf] hard-wrap da failed (%s) → latin-1 replace + küçük font", we2
                    )
                    # son emniyet: küçük font + '?' ile değiştirme (500 atmasın diye)
                    try:
                        current_family = pdf.font_family or "Uni"
                        current_size = getattr(pdf, "font_size_pt", 12) or 12
                        pdf.set_font(current_family, size=max(8, current_size - 2))
                    except Exception:
                        pass
                    safe = text.encode("latin-1", "replace").decode("latin-1")
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(eff_w, 10, safe)

            written += 1
        logger.info("[/word-to-pdf] paragraphs_written=%d", written)

        # 5) Geçici dosyaya kaydet
        temp_dir = tempfile.gettempdir()
        filename = f"converted_{uuid.uuid4().hex}.pdf"
        filepath = os.path.join(temp_dir, filename)
        logger.info("[/word-to-pdf] saving PDF to temp path=%s", filepath)
        pdf.output(filepath)
        try:
            size_bytes = os.path.getsize(filepath)
        except OSError:
            size_bytes = -1
        logger.info("[/word-to-pdf] PDF saved size_bytes=%d", size_bytes)

        # 6) Firebase Storage'a yükle
        bucket = storage.bucket()
        logger.info("[/word-to-pdf] firebase bucket=%s", getattr(bucket, "name", "<unknown>"))
        blob_path = f"converted_pdfs/{filename}"
        blob = bucket.blob(blob_path)
        logger.info("[/word-to-pdf] uploading to storage path=%s", blob_path)
        blob.upload_from_filename(filepath)
        logger.info("[/word-to-pdf] upload complete")
        blob.make_public()
        logger.info("[/word-to-pdf] made public url=%s", blob.public_url)

        # 7) Yanıt
        logger.info("[/word-to-pdf] END success")
        return {"status": "success", "file_url": blob.public_url}

    except Exception as e:
        logger.exception("[/word-to-pdf] ERROR: %s", e)
        raise HTTPException(status_code=500, detail=str(e))








