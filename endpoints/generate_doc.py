import os
import tempfile
import uuid
from fastapi import HTTPException
from main import app, client, DEFAULT_MODEL, DocRequest, storage
from docx import Document


@app.post("/generate-doc")
async def generate_doc(data: DocRequest):
    print("[/generate-doc] 📝 İstek alındı.")
    try:
        # 1. GPT'den içerik al
        print("[/generate-doc] 🧠 GPT'den içerik alınıyor...")
        completion = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[{"role": "user", "content": data.prompt}],
            max_tokens=1500
        )
        generated_text = completion.choices[0].message.content.strip()
        print("[/generate-doc] ✅ GPT içeriği alındı, uzunluk:", len(generated_text))
        print("[/generate-doc] 🔍 İlk 300 karakter:\n", generated_text[:300])

        # 2. Word belgesi oluştur
        print("[/generate-doc] 📄 Word belgesi oluşturuluyor...")
        doc = Document()
        doc.add_heading('Avenia Belgesi', 0)
        for i, paragraph in enumerate(generated_text.split("\n")):
            cleaned = paragraph.strip()
            if cleaned:
                doc.add_paragraph(cleaned)
                print(f"[/generate-doc] ➕ Paragraf {i+1}: {cleaned[:100]}")

        # 3. Geçici dosyaya kaydet
        temp_path = tempfile.gettempdir()
        filename = f"generated_{uuid.uuid4().hex}.docx"
        filepath = os.path.join(temp_path, filename)
        doc.save(filepath)
        print("[/generate-doc] 💾 Word dosyası kaydedildi:", filepath)

        # 4. Firebase Storage’a yükle
        print("[/generate-doc] ☁️ Firebase Storage’a yükleniyor...")
        bucket = storage.bucket()
        blob = bucket.blob(f"generated_docs/{filename}")
        blob.upload_from_filename(filepath)
        blob.make_public()
        print("[/generate-doc] 📤 Yükleme başarılı, link:", blob.public_url)

        # 5. URL’i dön
        return {
            "status": "success",
            "file_url": blob.public_url
        }

    except Exception as e:
        print("[/generate-doc] ❌ Hata oluştu:", str(e))
        raise HTTPException(status_code=500, detail=str(e))
