import json
import logging
import os
import tempfile
import uuid
import time
from typing import Any, Dict, List, Optional

import httpx
from fastapi import APIRouter, HTTPException, Request
from docx import Document
from firebase_admin import storage

from core.gemini_prompt import build_system_message, merge_parts_with_system
from core.firebase import db
from schemas import DocRequest
from endpoints.logging.utils_logging import log_gemini_request, log_gemini_response
from usage_tracking import build_base_event, finalize_event, parse_gemini_usage, enqueue_usage_update

logger = logging.getLogger("pdf_read_refresh.endpoints.generate_doc")
router = APIRouter()

SYSTEM_INSTRUCTION = (
    "You are a professional document preparation agent. "
    'Return ONLY valid JSON following this schema: {"title": "...", "sections": [{"heading": "...", "content": "..."}]}. '
    "Each section must have a concise heading and a paragraph-style content. "
    "Do not include markdown, code fences, or extra text outside JSON."
)


def _effective_model() -> str:
    model = os.getenv("GEMINI_DOC_MODEL") or os.getenv("GEMINI_SEARCH_MODEL") or "models/gemini-2.5-pro"
    if not model.startswith("models/"):
        model = f"models/{model}"
    return model


def _build_usage_context(request: Request, user_id: str, model: str) -> Optional[Dict[str, Any]]:
    if not user_id:
        return None
    token_payload = getattr(request.state, "token_payload", {}) or {}
    request_id = (
        request.headers.get("x-request-id")
        or request.headers.get("x-requestid")
        or f"req_{uuid.uuid4().hex}"
    )
    return build_base_event(
        request_id=request_id,
        user_id=user_id,
        endpoint="generate_doc",
        provider="gemini",
        model=model,
        token_payload=token_payload,
        request=request,
    )


def _enqueue_usage_event(
    usage_context: Optional[Dict[str, Any]],
    usage_data: Dict[str, int],
    latency_ms: int,
    *,
    status: str,
    error_code: Optional[str],
) -> None:
    if not usage_context or not db:
        return
    try:
        event = finalize_event(
            usage_context,
            input_tokens=usage_data.get("inputTokens", 0),
            output_tokens=usage_data.get("outputTokens", 0),
            latency_ms=latency_ms,
            status=status,
            error_code=error_code,
        )
        enqueue_usage_update(db, event)
    except Exception:
        logger.warning("Usage tracking failed for generate_doc", exc_info=True)


async def _call_gemini_json(prompt: str, system_message: Optional[str]) -> tuple[Dict[str, Any], Dict[str, Any]]:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail={"success": False, "error": "gemini_api_key_missing", "message": "GEMINI_API_KEY is required"},
        )

    model = _effective_model()
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"

    parts = [
        {"text": SYSTEM_INSTRUCTION},
        {"text": f"USER_PROMPT: {prompt}"},
    ]
    effective_parts = merge_parts_with_system(parts, system_message)
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": effective_parts,
            }
        ],
        "tools": [{"google_search": {}}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "temperature": 0.4,
            "maxOutputTokens": 2048,
        },
    }

    log_gemini_request(
        logger,
        "generate_doc",
        url=url,
        payload=payload,
        model=model,
    )
    logger.info("Gemini doc JSON request", extra={"model": model, "prompt_preview": prompt[:200]})

    async with httpx.AsyncClient(timeout=90) as client:
        resp = await client.post(url, json=payload)

    body_preview = (resp.text or "")[:800]
    response_json = resp.json() if resp.text else {}
    log_gemini_response(
        logger,
        "generate_doc",
        url=url,
        status_code=resp.status_code,
        response=response_json,
    )
    logger.info("Gemini doc JSON response", extra={"status": resp.status_code, "body_preview": body_preview})

    if not resp.ok:
        raise HTTPException(
            status_code=resp.status_code,
            detail={"success": False, "error": "gemini_doc_failed", "message": body_preview},
        )

    data = response_json
    parts = (data.get("candidates") or [{}])[0].get("content", {}).get("parts", [])
    if not parts:
        raise HTTPException(
            status_code=500,
            detail={"success": False, "error": "empty_response", "message": "Gemini returned no content"},
        )
    try:
        parsed = json.loads(parts[0].get("text") or "{}")
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=500,
            detail={"success": False, "error": "invalid_json", "message": "Gemini returned non-JSON content"},
        )
    return parsed, response_json


def _build_doc_from_json(doc_data: Dict[str, Any]) -> str:
    title = doc_data.get("title") or "Avenia Belgesi"
    sections: List[Dict[str, Any]] = doc_data.get("sections") or []

    doc = Document()
    doc.add_heading(title, 0)

    for idx, section in enumerate(sections):
        heading = (section or {}).get("heading") or f"Section {idx + 1}"
        content = (section or {}).get("content") or ""
        doc.add_heading(heading, level=1)
        for paragraph in str(content).split("\n"):
            clean = paragraph.strip()
            if clean:
                doc.add_paragraph(clean)

    temp_dir = tempfile.gettempdir()
    filename = f"generated_{uuid.uuid4().hex}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    return filepath, filename


@router.post("/generate-doc")
async def generate_doc(data: DocRequest, request: Request):
    logger.info("Generate doc request received", extra={"prompt_length": len(data.prompt or "")})
    payload = getattr(request.state, "token_payload", {}) or {}
    user_id = payload.get("uid") or payload.get("userId") or payload.get("sub") or ""
    usage_context: Optional[Dict[str, Any]] = None
    usage_data: Dict[str, int] = {}
    start_time = time.monotonic()
    status = "success"
    error_code = None
    try:
        # 1) Gemini'den yapılandırılmış JSON al
        model = _effective_model()
        usage_context = _build_usage_context(request, user_id, model)
        system_message = build_system_message(
            language=None,
            tone_key=data.tone_key,
            response_style=None,
            include_response_style=False,
            include_followup=False,
        )
        doc_json, response_json = await _call_gemini_json(data.prompt, system_message)
        usage_data = parse_gemini_usage(response_json)
        logger.debug("Gemini JSON parsed", extra={"keys": list(doc_json.keys())})

        # 2) Word dosyasını oluştur
        filepath, filename = _build_doc_from_json(doc_json)
        logger.info("Word document built", extra={"filepath": filepath})

        # 3) Firebase'e yükle
        bucket = storage.bucket()
        blob = bucket.blob(f"generated_docs/{filename}")
        blob.upload_from_filename(filepath)
        blob.make_public()
        logger.info("Firebase upload completed", extra={"file_url": blob.public_url})

        response_payload = {
            "status": "success",
            "file_url": blob.public_url,
            "title": doc_json.get("title"),
            "section_count": len(doc_json.get("sections") or []),
        }
        logger.debug("Generate doc response payload", extra={"response": response_payload})
        return response_payload

    except HTTPException as exc:
        status = "error"
        if isinstance(exc.detail, dict):
            error_code = exc.detail.get("error")
        error_code = error_code or "generate_doc_failed"
        logger.exception("Generate doc failed (HTTPException)")
        raise
    except Exception as e:
        status = "error"
        error_code = "generate_doc_failed"
        logger.exception("Generate doc failed")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        _enqueue_usage_event(
            usage_context,
            usage_data,
            int((time.monotonic() - start_time) * 1000),
            status=status,
            error_code=error_code,
        )

