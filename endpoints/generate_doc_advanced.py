# endpoints/generate_doc_advanced.py
import logging
import os
import re
import uuid
import math
import tempfile
from typing import List, Optional, Dict

from fastapi import HTTPException
from pydantic import BaseModel, Field

from main import app, client, DEFAULT_MODEL, storage  # mevcut import d√ºzeninizle aynƒ±
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger("pdf_read_refresh.endpoints.generate_doc_advanced")


# ------------ Pydantic Model (GELƒ∞≈ûMƒ∞≈û) ------------
class MarginsMM(BaseModel):
    top: Optional[float] = None
    bottom: Optional[float] = None
    left: Optional[float] = None
    right: Optional[float] = None


class DocAdvancedRequest(BaseModel):
    # ZORUNLU
    prompt: str = Field(..., description="Belgenin konusu/isteƒüi")

    # OPSƒ∞YONELLER
    language: Optional[str] = Field(None, description="tr/en, varsayƒ±lan: tr")
    title: Optional[str] = None
    page_goal: Optional[int] = Field(None, description="Yakla≈üƒ±k hedef sayfa")
    include_cover: Optional[bool] = True
    include_toc: Optional[bool] = True
    header_text: Optional[str] = None
    footer_text: Optional[str] = None
    page_numbers: Optional[bool] = True
    paper_size: Optional[str] = Field("A4", description="A4/Letter")
    orientation: Optional[str] = Field("portrait", description="portrait/landscape")
    margins_mm: Optional[MarginsMM] = None
    font: Optional[str] = "Calibri"
    font_size_pt: Optional[float] = 11
    line_spacing: Optional[float] = 1.15
    outline: Optional[List[str]] = None
    references: Optional[List[str]] = None
    reference_style: Optional[str] = Field(None, description="APA/MLA/IEEE")
    watermark_text: Optional[str] = None


# ------------ Yardƒ±mcƒ±lar ------------
def _set_page_setup(section, paper_size: str, orientation: str, margins: Optional[MarginsMM]):
    print("[/generate-doc-advanced] üìê Sayfa ayarlarƒ± uygulanƒ±yor...")
    # Kaƒüƒ±t boyutu
    if (paper_size or "").upper() == "LETTER":
        width, height = Inches(8.5), Inches(11)
    else:
        # A4: 210 √ó 297 mm
        width, height = Mm(210), Mm(297)

    # Y√∂n
    if (orientation or "").lower() == "landscape":
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = height, width
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width, section.page_height = width, height

    # Kenar bo≈üluklarƒ±
    if margins:
        if margins.top is not None:
            section.top_margin = Mm(margins.top)
        if margins.bottom is not None:
            section.bottom_margin = Mm(margins.bottom)
        if margins.left is not None:
            section.left_margin = Mm(margins.left)
        if margins.right is not None:
            section.right_margin = Mm(margins.right)

    print("[/generate-doc-advanced] ‚úÖ Sayfa ayarlarƒ± tamam.")


def _set_document_style(doc: Document, font_name: str, font_size_pt: float, line_spacing: float):
    print("[/generate-doc-advanced] üñãÔ∏è Stil ayarlarƒ± uygulanƒ±yor...")
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)

    pf = style.paragraph_format
    try:
        pf.line_spacing = line_spacing  # multiple spacing (√∂rn: 1.15)
    except Exception:
        pass

    print("[/generate-doc-advanced] ‚úÖ Stil ayarlarƒ± tamam.")


def _add_page_number(footer_para):
    print("[/generate-doc-advanced] üî¢ Sayfa numarasƒ± alanƒ± ekleniyor...")
    run = footer_para.add_run()

    fldChar1 = OxmlElement('w:fldChar')     # begin
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '

    fldChar2 = OxmlElement('w:fldChar')     # separate
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')     # end
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    print("[/generate-doc-advanced] ‚úÖ Sayfa numarasƒ± eklendi.")


def _ensure_header_footer(doc: Document, header_text: Optional[str], footer_text: Optional[str], page_numbers: bool):
    print("[/generate-doc-advanced] üß¢ Header & Footer ayarlanƒ±yor...")
    section = doc.sections[0]

    if header_text:
        hp = section.header.paragraphs[0]
        hp.text = header_text
        hp.style = doc.styles['Header']
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fp = section.footer.paragraphs[0]
    if footer_text:
        fp.text = footer_text + "   "
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if page_numbers:
        _add_page_number(fp)

    print("[/generate-doc-advanced] ‚úÖ Header & Footer tamam.")


def _add_watermark_like_header(doc: Document, watermark_text: str):
    print("[/generate-doc-advanced] üíß Filigran (basit) ekleniyor...")
    section = doc.sections[0]
    p = section.header.add_paragraph()
    run = p.add_run(watermark_text.upper())
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("[/generate-doc-advanced] ‚úÖ Filigran eklendi.")


def _insert_toc(paragraph):
    print("[/generate-doc-advanced] üìö TOC alanƒ± ekleniyor...")
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')  # begin
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar2 = OxmlElement('w:fldChar')  # separate
    fldChar2.set(qn('w:fldCharType'), 'separate')

    run_after = paragraph.add_run("ƒ∞√ßindekiler i√ßin Word'de: Saƒü tƒ±k ‚Üí Alanlarƒ± G√ºncelle")
    run_after.italic = True

    fldChar3 = OxmlElement('w:fldChar')  # end
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    print("[/generate-doc-advanced] ‚úÖ TOC eklendi.")


def _add_cover_page(doc: Document, title: str, subtitle: Optional[str]):
    print("[/generate-doc-advanced] üßæ Kapak sayfasƒ± ekleniyor...")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)

    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(14)

    doc.add_page_break()
    print("[/generate-doc-advanced] ‚úÖ Kapak eklendi ve sayfa sonu konuldu.")


def _markdown_to_docx(doc: Document, text: str, line_spacing: float):
    """
    Basit markdown: '## ' ‚Üí Heading 1, '### ' ‚Üí Heading 2
    Bo≈ü satƒ±rlarla paragraflarƒ± ayƒ±rƒ±r.
    """
    print("[/generate-doc-advanced] üß© Markdown i√ßeriƒüi Word'e i≈üleniyor...")
    lines = text.splitlines()
    buf = []

    def flush_paragraph():
        if not buf:
            return
        para_text = " ".join(buf).strip()
        if para_text:
            p = doc.add_paragraph(para_text)
            try:
                p.paragraph_format.line_spacing = line_spacing
            except Exception:
                pass
        buf.clear()

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:].strip(), level=2)
            continue

        if line.strip() == "":
            flush_paragraph()
        else:
            buf.append(line)

    flush_paragraph()
    print("[/generate-doc-advanced] ‚úÖ Markdown i≈üleme tamam.")


def _add_references(doc: Document, refs: List[str], line_spacing: float, style_name="Heading 1"):
    print("[/generate-doc-advanced] üîó Kaynak√ßa b√∂l√ºm√º ekleniyor...")
    doc.add_page_break()
    doc.add_heading("Kaynak√ßa", level=1)
    for ref in refs:
        p = doc.add_paragraph(ref)
        try:
            p.paragraph_format.line_spacing = line_spacing
        except Exception:
            pass
    print("[/generate-doc-advanced] ‚úÖ Kaynak√ßa eklendi.")


def _estimate_words_per_page(
    paper_size: str,
    orientation: str,
    margins: Optional[MarginsMM],
    font_size_pt: float,
    line_spacing: float
) -> int:
    """
    Heuristik WPP (words per page) tahmini.
    Baseline: A4, dikey, 20mm kenar, Calibri 11pt, 1.15 -> ~350 kelime/sayfa.
    """
    print("[/generate-doc-advanced] üìè WPP hesaplanƒ±yor...")
    base_wpp = 350.0

    # Kaƒüƒ±t √∂l√ß√ºs√º (mm)
    if (paper_size or "").upper() == "LETTER":
        full_w, full_h = 215.9, 279.4  # 8.5x11 in√ß
    else:
        full_w, full_h = 210.0, 297.0  # A4

    if (orientation or "").lower() == "landscape":
        full_w, full_h = full_h, full_w

    # Kenar bo≈üluklarƒ± (mm)
    top = margins.top if (margins and margins.top is not None) else 20.0
    bottom = margins.bottom if (margins and margins.bottom is not None) else 20.0
    left = margins.left if (margins and margins.left is not None) else 20.0
    right = margins.right if (margins and margins.right is not None) else 20.0

    text_w = max(10.0, full_w - (left + right))
    text_h = max(10.0, full_h - (top + bottom))

    # Baseline alanƒ±: A4 dikey 20mm kenar: 170x257 mm
    base_area = 170.0 * 257.0
    area = text_w * text_h
    area_factor = area / base_area

    # Font ve satƒ±r aralƒ±ƒüƒ± etkisi (yakla≈üƒ±k)
    size_factor = 11.0 / float(font_size_pt or 11.0)
    spacing_factor = 1.15 / float(line_spacing or 1.15)

    wpp = base_wpp * area_factor * size_factor * spacing_factor
    wpp_clamped = int(max(180, min(700, wpp)))  # abartƒ±yƒ± kƒ±rp
    print(f"[/generate-doc-advanced] ‚úÖ WPP ‚âà {wpp_clamped} (area_factor={area_factor:.2f}, size={font_size_pt}, spacing={line_spacing})")
    return wpp_clamped


def _count_words(text: str) -> int:
    wc = len(re.findall(r"\w+", text, flags=re.UNICODE))
    return wc


def _build_system_and_user_prompts(data: DocAdvancedRequest):
    lang = (data.language or "tr").lower()

    target_words = None
    if data.page_goal and data.page_goal > 0:
        # _estimate_words_per_page ile endpoint i√ßinde daha iyi hesaplanacak;
        # burada yine de bir hedef g√∂nderiyoruz.
        target_words = int(data.page_goal * 350)

    system = (
        f"You are a senior technical writer. Write in {lang.upper()} language. "
        f"Structure content with '##' and '###' headings. "
        f"Use clear, concise, well-organized prose suitable for a formal handbook."
    )

    outline_str = ""
    if data.outline and len(data.outline) > 0:
        outline_str = "\n".join([f"- {h}" for h in data.outline])

    user = [
        f"TITLE: {data.title or 'Belge'}",
        f"PROMPT: {data.prompt}",
        f"TARGET_WORDS: {target_words or 'flexible'}",
        "OUTLINE:",
        outline_str if outline_str else "(propose a reasonable outline and use it)",
        "",
        "REQUIREMENTS:",
        "- Start each top-level section with '## ' + title",
        "- For subsections use '### '",
        "- Use paragraphs (no bullet walls)",
        "- Be specific and practical",
        "- Avoid filler content",
    ]
    user_text = "\n".join(user)
    return system, user_text


def _generate_text_to_word_budget(data: DocAdvancedRequest, system: str, base_user_text: str) -> str:
    """
    Hedef sayfa/kelime miktarƒ±na ula≈üana kadar iteratif i√ßerik √ºretir.
    Her turda 'yeni en az N kelime' isteyerek devam ettirir.
    """
    print("[/generate-doc-advanced] üßÆ Kelime b√ºt√ßeli √ºretim ba≈ülƒ±yor...")
    # WPP tahmini
    wpp = _estimate_words_per_page(
        paper_size=data.paper_size or "A4",
        orientation=data.orientation or "portrait",
        margins=data.margins_mm,
        font_size_pt=float(data.font_size_pt or 11),
        line_spacing=float(data.line_spacing or 1.15),
    )

    # Hedef kelime
    if data.page_goal and data.page_goal > 0:
        target_words = int(max(1, data.page_goal) * wpp)
    else:
        # Sayfa hedefi yoksa geni≈ü bir aralƒ±k belirleyelim (tek tur da olabilir)
        target_words = 1200

    # G√ºvenlik limitleri
    target_words = int(min(max(target_words, 400), 25000))  # 400‚Äì25000 arasƒ±
    print(f"[/generate-doc-advanced] üéØ Hedef kelime: {target_words} (page_goal={data.page_goal}, wpp={wpp})")

    accumulated = ""
    passes = 0
    max_passes = 8  # fazla abartmadan
    remaining = target_words

    while remaining > 0 and passes < max_passes:
        # Her turda 800‚Äì1400 kelime civarƒ± isteyelim
        chunk_goal = int(min(remaining, 1200 if passes > 0 else 1400))
        chunk_goal = max(600, chunk_goal)  # √ßok k√º√ß√ºk olmasƒ±n

        if passes == 0:
            # ƒ∞lk tur: temel prompt
            user_msg = (
                base_user_text
                + "\n\n"
                + f"Write at least {chunk_goal} NEW words. "
                  "Cover the outline thoroughly with '##' and '###' headings. "
                  "Use paragraphs, not bullet walls. Avoid repetition."
            )
            messages = [
                {"role": "system", "content": system},
                {"role": "user", "content": user_msg},
            ]
        else:
            # Sonraki turlar: devam isteƒüi
            last_slice = accumulated[-6000:]  # modele biraz baƒülam verelim
            continue_msg = (
                f"CONTINUE the same document. Write at least {chunk_goal} NEW words. "
                f"Do not repeat previous sentences or headings. Deepen analysis, add examples, and expand remaining/underdeveloped sections. "
                f"Keep the same '##'/'###' structure and formal tone."
            )
            messages = [
                {"role": "system", "content": system},
                {"role": "assistant", "content": last_slice},
                {"role": "user", "content": continue_msg},
            ]

        print(f"[/generate-doc-advanced] üß™ Tur #{passes+1} isteniyor (chunk_goal‚âà{chunk_goal} kelime)...")
        completion = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=messages,
            temperature=0.7,
            max_tokens=min(8192, int(chunk_goal * 2))  # kaba d√∂n√º≈ü√ºm: ~2 token/kelime
        )
        part = completion.choices[0].message.content.strip()
        part_wc = _count_words(part)
        print(f"[/generate-doc-advanced] ‚úÖ Tur #{passes+1} alƒ±ndƒ±: ~{part_wc} kelime")
        print(f"[/generate-doc-advanced] üîç √ñnizleme (ilk 300):\n{part[:300]}")

        # Birle≈ütir
        if accumulated:
            accumulated += "\n\n" + part
        else:
            accumulated = part

        total_wc = _count_words(accumulated)
        remaining = max(0, target_words - total_wc)
        print(f"[/generate-doc-advanced] üìà Birikimli kelime: {total_wc} | Kalan hedef: {remaining}")

        # Eƒüer bir turda beklenenden √ßok az geldiyse, bir tur daha denemek mantƒ±klƒ± olabilir
        passes += 1

        # Erken √ßƒ±kƒ±≈ü: hedefin %90‚Äôƒ±na geldiysek yeterli kabul edelim
        if total_wc >= 0.9 * target_words:
            print("[/generate-doc-advanced] ‚úÖ Hedefin %90+‚Äôƒ±na ula≈üƒ±ldƒ±, √ºretim tamamlanƒ±yor.")
            break

    print("[/generate-doc-advanced] üèÅ Kelime b√ºt√ßeli √ºretim bitti.")
    return accumulated


# ------------ ENDPOINT ------------
@app.post("/generate-doc-advanced")
async def generate_doc_advanced(data: DocAdvancedRequest):
    logger.info(
        "Generate doc advanced request received",
        extra={"prompt_preview": (data.prompt or "")[:200], "page_goal": data.page_goal},
    )
    try:
        logger.info("Starting GPT content generation")
        system, user_text = _build_system_and_user_prompts(data)

        # Eƒüer page_goal varsa, iteratif kelime b√ºt√ßeli √ºretim yap
        if data.page_goal and int(data.page_goal) > 0:
            logger.info("page_goal detected - running iterative budgeted generation")
            generated = _generate_text_to_word_budget(data, system, user_text)
        else:
            logger.info("No page_goal - single-shot generation")
            completion = client.chat.completions.create(
                model=DEFAULT_MODEL,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_text},
                ],
                temperature=0.7,
                max_tokens=3500
            )
            generated = completion.choices[0].message.content.strip()

        logger.info("Generated text ready", extra={"char_length": len(generated)})
        logger.debug("Generated text preview", extra={"preview": generated[:400]})

        # 2) Word dok√ºmanƒ±nƒ± kur
        logger.info("Building Word document")
        doc = Document()

        # Stil & Sayfa ayarlarƒ±
        _set_document_style(doc, data.font or "Calibri", float(data.font_size_pt or 11), float(data.line_spacing or 1.15))
        _set_page_setup(doc.sections[0], data.paper_size or "A4", data.orientation or "portrait", data.margins_mm)

        if data.watermark_text:
            _add_watermark_like_header(doc, data.watermark_text)

        # Kapak
        if data.include_cover:
            _add_cover_page(doc, data.title or "Avenia Belgesi", data.prompt)

        # ƒ∞√ßindekiler
        if data.include_toc:
            doc.add_heading("ƒ∞√ßindekiler", level=1)
            _insert_toc(doc.add_paragraph())
            doc.add_page_break()

        # ƒ∞√ßerik (markdown basit d√∂n√º≈ü√ºm)
        _markdown_to_docx(doc, generated, float(data.line_spacing or 1.15))

        # Kaynak√ßa
        if data.references and len(data.references) > 0:
            _add_references(doc, data.references, float(data.line_spacing or 1.15))

        # Header/Footer & Sayfa No
        _ensure_header_footer(doc, data.header_text, data.footer_text, bool(data.page_numbers if data.page_numbers is not None else True))

        # 3) Ge√ßici dosyaya kaydet
        temp_dir = tempfile.gettempdir()
        filename = f"generated_{uuid.uuid4().hex}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        logger.info("Word document saved", extra={"filepath": filepath})

        # 4) Firebase Storage‚Äôa y√ºkle
        logger.info("Uploading document to Firebase Storage")
        bucket = storage.bucket()
        blob = bucket.blob(f"generated_docs/{filename}")
        blob.upload_from_filename(filepath)
        blob.make_public()
        logger.info("Firebase upload completed", extra={"file_url": blob.public_url})

        # 5) D√∂n√º≈ü
        return {
            "status": "success",
            "file_url": blob.public_url
        }

    except Exception as e:
        logger.exception("Generate doc advanced failed")
        raise HTTPException(status_code=500, detail=str(e))
